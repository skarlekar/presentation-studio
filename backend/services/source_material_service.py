"""
Source material service — extracts plain text from uploaded files.

Supported formats: PDF (.pdf), plain text (.txt), Markdown (.md), Word (.docx).
Maximum file size: 10 MB.

Usage:
    text = await extract_text_from_file(file_bytes, "report.pdf")
"""
import io
from pathlib import Path
from typing import Optional

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = frozenset({".pdf", ".txt", ".docx", ".md"})


class SourceMaterialError(Exception):
    """Raised when file extraction fails due to type, size, or content issues."""


async def extract_text_from_file(
    file_content: bytes,
    filename: str,
    content_type: Optional[str] = None,
) -> str:
    """Extract plain text from an uploaded file.

    Args:
        file_content: Raw file bytes.
        filename: Original filename — used to determine file type by extension.
        content_type: Optional MIME type hint (not relied upon for type detection).

    Returns:
        Extracted plain text string, stripped of leading/trailing whitespace.

    Raises:
        SourceMaterialError: If the file exceeds 10 MB, the type is unsupported,
                             or text extraction fails.
    """
    if len(file_content) > MAX_FILE_SIZE_BYTES:
        size_mb = len(file_content) / (1024 * 1024)
        raise SourceMaterialError(
            f"File size {size_mb:.1f} MB exceeds the 10 MB limit."
        )

    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise SourceMaterialError(
            f"Unsupported file type '{ext}'. "
            f"Allowed extensions: {', '.join(sorted(ALLOWED_EXTENSIONS))}."
        )

    if ext in (".txt", ".md"):
        return _extract_text(file_content)
    elif ext == ".pdf":
        return _extract_pdf(file_content)
    elif ext == ".docx":
        return _extract_docx(file_content)
    else:
        # Should never reach here given the allowlist check above
        raise SourceMaterialError(f"No extractor registered for extension '{ext}'.")


def _extract_text(content: bytes) -> str:
    """Decode a plain text or Markdown file, trying common encodings."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    raise SourceMaterialError(
        "Could not decode text file — encoding is unknown or not supported."
    )


def _extract_pdf(content: bytes) -> str:
    """Extract text from a PDF using pypdf.

    Raises:
        SourceMaterialError: If pypdf is not installed, the PDF is image-only,
                             or extraction otherwise fails.
    """
    try:
        from pypdf import PdfReader  # type: ignore[import]
    except ImportError:
        raise SourceMaterialError(
            "pypdf is not installed. Install it with: pip install pypdf"
        )

    try:
        reader = PdfReader(io.BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())

        extracted = "\n\n".join(pages)
        if not extracted.strip():
            raise SourceMaterialError(
                "PDF appears to be scanned or image-only — no text could be extracted. "
                "Please provide a text-based PDF or paste the content directly."
            )
        return extracted
    except SourceMaterialError:
        raise
    except Exception as exc:
        raise SourceMaterialError(f"PDF extraction failed: {exc}") from exc


def _extract_docx(content: bytes) -> str:
    """Extract text from a DOCX file using python-docx.

    Extracts paragraph text and table cell content.

    Raises:
        SourceMaterialError: If python-docx is not installed, the file is empty,
                             or extraction otherwise fails.
    """
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        raise SourceMaterialError(
            "python-docx is not installed. Install it with: pip install python-docx"
        )

    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        # Extract paragraph text
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in parts:
                        parts.append(text)

        extracted = "\n\n".join(parts)
        if not extracted.strip():
            raise SourceMaterialError(
                "DOCX file appears to be empty — no text content was found."
            )
        return extracted
    except SourceMaterialError:
        raise
    except Exception as exc:
        raise SourceMaterialError(f"DOCX extraction failed: {exc}") from exc
